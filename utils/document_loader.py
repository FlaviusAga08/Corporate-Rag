import os
import re
from collections import Counter
from typing import List, Dict, Any

import fitz  # pymupdf
from docx import Document

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

def _chunk_text(text: str, filename: str, page_num: int) -> List[Dict[str, Any]]:
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks: List[Dict[str, Any]] = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 <= CHUNK_SIZE:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append({
                    "text": current,
                    "metadata": {"source": filename, "page": page_num}
                })
            if len(para) > CHUNK_SIZE:
                start = 0
                while start < len(para):
                    chunks.append({
                        "text": para[start:start + CHUNK_SIZE].strip(),
                        "metadata": {"source": filename, "page": page_num}
                    })
                    start += CHUNK_SIZE - CHUNK_OVERLAP
                current = ""
            else:
                current = para

    if current:
        chunks.append({
            "text": current,
            "metadata": {"source": filename, "page": page_num}
        })

    return [c for c in chunks if len(c["text"]) > 30]


def _strip_pdf_noise(pages_lines: List[List[str]], n_pages: int) -> List[List[str]]:
    counts = Counter(
        line.strip()
        for lines in pages_lines
        for line in lines
        if line.strip()
    )
    noise = {
        line for line, count in counts.items()
        if count >= max(2, n_pages * 0.4) and len(line) < 120
    }
    
    cleaned_pages = []
    for lines in pages_lines:
        cleaned_lines = [l for l in lines if l.strip() not in noise]
        cleaned_pages.append(cleaned_lines)
    return cleaned_pages


def load_pdf(file_path: str) -> List[Dict[str, Any]]:
    doc = fitz.open(file_path)
    filename = os.path.basename(file_path)
    pages_lines = [page.get_text().split("\n") for page in doc]
    cleaned_pages_lines = _strip_pdf_noise(pages_lines, len(doc))
    
    pdf_chunks = []
    for page_idx, lines in enumerate(cleaned_pages_lines):
        page_text = "\n".join(lines)
        # Procesăm chunk-urile salvând pagina curentă (1-indexed)
        pdf_chunks.extend(_chunk_text(page_text, filename, page_idx + 1))
    return pdf_chunks


def load_docx(file_path: str) -> List[Dict[str, Any]]:
    doc = Document(file_path)
    filename = os.path.basename(file_path)
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Fisierele Word nu au pagini fixe în mod nativ, punem pagina 1 generic
    return _chunk_text("\n\n".join(parts), filename, 1)


def load_documents(directory: str) -> List[Dict[str, Any]]:
    all_chunks = []
    for filename in sorted(os.listdir(directory)):
        path = os.path.join(directory, filename)
        if filename.endswith(".pdf"):
            all_chunks.extend(load_pdf(path))
        elif filename.endswith(".docx"):
            all_chunks.extend(load_docx(path))
    return all_chunks